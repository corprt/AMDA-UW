from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import date
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.table import WD_ALIGN_VERTICAL

from docx.oxml import OxmlElement, ns
from docx.oxml.ns import qn

from docx.shared import Inches
from docx.opc.constants import RELATIONSHIP_TYPE
import pandas as pd
from docx.shared import RGBColor
from docx.enum.style import WD_STYLE_TYPE
import os
import streamlit as st
write_output_path = './output/'


# Add header footer to document
def add_header(doc, header_text):
    #doc = Document()
    section = doc.sections[0]
    header = section.header
    p = header.paragraphs[0]
    p.clear()  # Clear existing content
    p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    run1 = p.add_run(header_text)
    run1.font.size = Pt(12)
    run1.font.italic = True

def add_footer(doc, footer_text):
    #doc = Document()
    section = doc.sections[0]
    footer = section.footer
    p = footer.paragraphs[0]
    p.clear()  # Clear existing content
    p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    run1 = p.add_run(footer_text)
    run1.font.size = Pt(8)
    run1.font.italic = True

#===================================================================================================================================================================#

# Document section creation using heading, description, tables, observations
def create_doc_section(document, document_title, heading, description, tables, observation, sec_no, byline):
    
    if sec_no > 0:
        document.add_page_break()

    add_header(document, document_title) # header
    add_footer(document, "Confidential") # footer

    # heading and paragraph
    if heading is not None:
        document.add_heading(heading,1)        
        
        if byline != "NA":
            b = document.add_paragraph(style='CommentsStyle')
            run = b.add_run(byline)
            font = run.font
            font.size = Pt(12)
            #font.name = 'Arial'
            font.italic = True
            font.color.rgb = RGBColor(46, 77, 167)
            b.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            #document.add_paragraph('\n')
            
        p = document.add_paragraph(description, style = 'CommentsStyle')
#         document.add_paragraph('\n')

    #Processing tables
    if tables is not None:
        doc_table = document.add_table(rows=1, cols=len(tables.columns))
        doc_table.style = 'Table Grid'
        header_cells = doc_table.rows[0].cells
        # Populate table headers
        for i, column_name in enumerate(tables.columns):
            header_cells[i].text = column_name
            run = header_cells[i].paragraphs[0].runs[0]
            run.bold = True
            run.font.size = Pt(12)   
        # Add the table rows
        for _, row in tables.iterrows():
            row_cells = doc_table.add_row().cells
            for i, value in enumerate(row.values):
                row_cells[i].text = str(value)
   
    # Add observation 
    if observation is not None:
        document.add_heading("Observation of test", 2)
        d = document.add_paragraph(observation, style = 'CommentsStyle')
        document.add_paragraph('\n')

    if tables is not None:
        document.add_paragraph('\n')


#=========================================================================================================================================#
def create_doc_subsection_data_overview(document, document_title,
               subsection, subsection_description, byline, subsections_tables_list, subsections_observations_dict,
                print_charts,print_tables, table_names = [] ):
    
    add_header(document, document_title) # header
    add_footer(document, "Confidential") # footer

    #===============writing subsections====================#

    # subsections, descriptions of subsections
    document.add_heading(subsection,2)
    if byline != "NA":
            b = document.add_paragraph(style = 'CommentsStyle')
            run = b.add_run(byline)
            font = run.font
            font.size = Pt(12)
            #font.name = 'Arial'
            font.italic = True
            font.color.rgb = RGBColor(46, 77, 167)
            b.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            #document.add_paragraph('\n')
            

        #Processing tables for subsection
    if print_tables: # if table printing is needed
        for subsec_tables in subsections_tables_list:
            subsec_doc_table = document.add_table(rows=1, cols=len(subsec_tables.columns))
            subsec_doc_table.style = 'Table Grid'
            header_cells = subsec_doc_table.rows[0].cells
            # Populate table headers
            for i, column_name in enumerate(subsec_tables.columns):
                header_cells[i].text = column_name
                run = header_cells[i].paragraphs[0].runs[0]
                run.bold = True
                run.font.size = Pt(12)   
            # Add the table rows
            for _, row in subsec_tables.iterrows():
                row_cells = subsec_doc_table.add_row().cells
                for i, value in enumerate(row.values):
                    row_cells[i].text = str(value)
            
            if subsections_tables_list != []: # if any table printed
                document.add_paragraph('\n')
    else:
        # add placeholders
        for tn in table_names:
            document.add_paragraph('Add excel file' + str(tn), style = 'CommentsStyle')
        #check if subsection description is dictionary for hard coded input
    if isinstance(subsection_description, dict):
        for key, value in subsection_description.items():
            heading = document.add_paragraph(style = 'CommentsStyle')
            run = heading.add_run(key + ":")
            font = run.font
            font.bold = True
            description = document.add_paragraph(value, style = 'CommentsStyle')
            
    else:
        p = document.add_paragraph(subsection_description, style = 'CommentsStyle')
        
    document.add_paragraph('\n')
    if print_charts:
        subsec = subsection.split('.')[2].strip() 
        image_full_path = os.path.join('output', 'chartimages', str(subsec)+'.png')
        document.add_picture(image_full_path, width=Inches(6))
        document.add_page_break()
#     st.write(subsections_observations_dict["Data Overview2"])
    document.add_paragraph(subsections_observations_dict["Data Overview2"])

#=========================================================================================================================================#


# Document creation using heading, description, tables, observations
def create_doc_subsection(document, document_title,
               subsection, subsection_description, byline, subsections_tables_list, subsections_observations_dict,
                print_charts,print_tables, table_names = [] ):

    add_header(document, document_title) # header
    add_footer(document, "Confidential") # footer

    #===============writing subsections====================#

    # subsections, descriptions of subsections
    document.add_heading(subsection,2)
    if byline != "NA":
            b = document.add_paragraph(style = 'CommentsStyle')
            run = b.add_run(byline)
            font = run.font
            font.size = Pt(12)
            #font.name = 'Arial'
            font.italic = True
            font.color.rgb = RGBColor(46, 77, 167)
            b.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            #document.add_paragraph('\n')
            
    #check if subsection description is dictionary for hard coded input
    if isinstance(subsection_description, dict):
        for key, value in subsection_description.items():
            heading = document.add_paragraph(style = 'CommentsStyle')
            run = heading.add_run(key + ":")
            font = run.font
            font.bold = True
            description = document.add_paragraph(value, style = 'CommentsStyle')
            
    else:
        p = document.add_paragraph(subsection_description, style = 'CommentsStyle')
        
    document.add_paragraph('\n')

    #Processing tables for subsection
    if print_tables: # if table printing is needed
        for subsec_tables in subsections_tables_list:
            subsec_doc_table = document.add_table(rows=1, cols=len(subsec_tables.columns))
            subsec_doc_table.style = 'Table Grid'
            header_cells = subsec_doc_table.rows[0].cells
            # Populate table headers
            for i, column_name in enumerate(subsec_tables.columns):
                header_cells[i].text = column_name
                run = header_cells[i].paragraphs[0].runs[0]
                run.bold = True
                run.font.size = Pt(12)   
            # Add the table rows
            for _, row in subsec_tables.iterrows():
                row_cells = subsec_doc_table.add_row().cells
                for i, value in enumerate(row.values):
                    row_cells[i].text = str(value)
            
            if subsections_tables_list != []: # if any table printed
                document.add_paragraph('\n')
    else:
        # add placeholders
        for tn in table_names:
            document.add_paragraph('Add excel file' + str(tn), style = 'CommentsStyle')
    if print_charts:
        subsec = subsection.split('.')[2].strip() 
        image_full_path = os.path.join('output', 'chartimages', str(subsec)+'.png')
        document.add_picture(image_full_path, width=Inches(6))
        document.add_page_break()



def create_doc_testing_plan(document, document_title, subsec, test_name, subsections_observations_list, byline):
    add_header(document, document_title) # header
    add_footer(document, "Confidential") # footer
    document.add_heading(subsec, 2)
    if byline != "NA":
            b = document.add_paragraph(style = 'CommentsStyle')
            run = b.add_run(byline)
            font = run.font
            font.size = Pt(12)
            #font.name = 'Arial'
            font.italic = True
            font.color.rgb = RGBColor(46, 77, 167)
            b.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            #document.add_paragraph('\n')
    
    for j, test in enumerate(test_name):
        document.add_heading("5.1." +str(j+1) + " " +str(test) ,3) # add test heading
        observation_text = subsections_observations_list[j]
        document.add_paragraph(observation_text, style = 'CommentsStyle')
#=============================================================================#

def create_doc_subsection_test_results(document, subsec, byline, subsections_tables_list, subsections_observations_list,
               test_names_list, chart_ot):
    
    document.add_heading(subsec,2)
    if byline != "NA":
            b = document.add_paragraph(style = 'CommentsStyle')
            run = b.add_run(byline)
            font = run.font
            font.size = Pt(12)
            #font.name = 'Arial'
            font.italic = True
            font.color.rgb = RGBColor(46, 77, 167)
            b.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            #document.add_paragraph('\n')
    
    # add sub sub sections for each test
    for j, tst2 in enumerate(test_names_list):
        document.add_heading("5.2."+str(j+1)+" Result of " + str(tst2) + " test",3) # add test heading

        # print tables
        dict_index = {0:[0,1], 1:[2,3], 2:[4,5], 3:[6,7], 4:[8,9]} # map the index of test name with index of corresponding table in table list
        # if table for test is not present in 
        tables_list = subsections_tables_list[dict_index[j][0]:dict_index[j][1] + 1]

        if tst2 == "GINI":
            concatenated = []
            concatenated.append(pd.concat([tables_list[0], tables_list[1]]))
            tables_list = concatenated
            
        elif tst2 == "AUC":
            concatenated = []
            concatenated.append(pd.concat([tables_list[0], tables_list[1]]))
            tables_list = concatenated

        elif tst2 == "RMSE":
            concatenated = []
            concatenated.append(pd.concat([tables_list[0], tables_list[1]]))
            tables_list = concatenated


        # elif tst2 == "Rank Ordering":
        #     concatenated = []
        #     concatenated.append(pd.concat([tables_list[0], tables_list[1]]))
        #     tables_list = concatenated
            
        for subsec_tables in tables_list:

            # special handling for KS results => drop columns to save space
            if tst2 == "KS Statistic":
                subsec_tables = subsec_tables.reset_index()
                subsec_tables = subsec_tables.loc[:,['Decile', 'min_prob', 'max_prob', 
                                                     'cum_event_capture_rate', 'cum_nonevent_capture_rate', 'KS']]
                subsec_tables.rename(columns = {'min_prob' : 'Min Probability', 'max_prob' : 'Max Probability',
                                                'cum_event_capture_rate' : 'Cumulative Event Capture Rate',
                                                'cum_nonevent_capture_rate' : 'Cumulative Non Event Capture Rate'}, inplace = True)
                

            subsec_doc_table = document.add_table(rows=1, cols=len(subsec_tables.columns))
            subsec_doc_table.style = 'Table Grid'
            header_cells = subsec_doc_table.rows[0].cells
            # Populate table headers
            for i, column_name in enumerate(subsec_tables.columns):
                header_cells[i].text = column_name
                run = header_cells[i].paragraphs[0].runs[0]
                run.bold = True
                run.font.size = Pt(12)   
            # Add the table rows
            for _, row in subsec_tables.iterrows():
                row_cells = subsec_doc_table.add_row().cells
                for i, value in enumerate(row.values):
                    row_cells[i].text = str(value)
            document.add_paragraph('\n')
        # special handling for KS test => refer to excel in output folder for complete picture

        if tst2 == "KS Statistic":
            document.add_paragraph("*(For full data refer to KS Statistic.xlsx in output folder)\n", style = 'CommentsStyle')

        observation_text = subsections_observations_list[j]
        document.add_paragraph(observation_text, style = 'CommentsStyle')
        #==================SC=======================#
        #Block for adding png charts for each tests
        if tst2 in chart_ot:
            image = str(tst2)+".png"
            # document.add_paragraph("Results of the" + str(tst2) +"test")
            image_full_path = os.path.join('output', 'chartimages', image)
            document.add_picture(image_full_path, width=Inches(6))
            document.add_page_break()
#===========================================================================#

def add_intro(document, content):
    section = document.sections[0]  # Get the first section of the document
    new_paragraph = document.add_paragraph(content, style='Normal')
    section.start_type   # Set the start type of the section to a new page
    new_paragraph.clear()
    new_run = new_paragraph.add_run(content)
    new_run.bold = True
    return document
#=============================================================================#

def add_toc(doc, sections_dict):
    toc = doc.add_paragraph()
    toc.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    toc_run = toc.add_run('Table of Contents')
    toc_run.bold = True
    toc_run.font.size = Pt(12)

    for section, subsections in sections_dict.items():
        # Add section
        section_para = doc.add_paragraph(style = 'CommentsStyle')
        section_run = section_para.add_run(section)
        section_run.bold = True
        section_run.font.size = Pt(10)
        section_para.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

        for subsection in subsections:
            # Add subsections with indentation
            subsection_para = doc.add_paragraph(style = 'CommentsStyle')
            subsection_run = subsection_para.add_run(subsection)
            # Adjust indentation for subsections
            subsection_para.paragraph_format.left_indent = Pt(20)
            subsection_run.font.size = Pt(10)
            subsection_para.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    doc.add_page_break()  # Add a page break after the table of contents

#============================================================================#

def add_titlepage(document, document_title, model_name):
    title = document.add_heading(document_title,0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.runs[0].font.size = Pt(30)
    title.runs[0].font.name = 'Arial'
    title.runs[0].bold = True

    model = document.add_paragraph(model_name)
    model.alignment = WD_ALIGN_PARAGRAPH.CENTER
    model.runs[0].font.size = Pt(24)
    model.runs[0].font.name = 'Arial'
    model.runs[0].bold = True
    
    today = date.today().strftime("%B %d,%Y")
    date_para = document.add_paragraph(today)
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_para.runs[0].font.size = Pt(14)
    date_para.runs[0].font.name = 'Arial'

    document.add_page_break()

# def add_hyperlink_to_document(doc, url):
#     """
#     Add a hyperlink to the given document.
#     """
#     # Create a new centered paragraph
#     para = doc.add_paragraph()
#     para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

#     # Add text and hyperlink to the paragraph
#     text = "Click here to open Excel file"
#     run = para.add_run(text)
#     run.font.size = Pt(12)
#     run.font.color.rgb = RGBColor(0, 0, 255)  # Blue color
#     run.font.underline = True

#     # Create a hyperlink element
#     hyperlink = OxmlElement('w:hyperlink')
#     hyperlink.set('{http://www.w3.org/1999/xlink}href', url)

#     # Insert the hyperlink element into the run
#     run._r.append(hyperlink)
def create_element(name):
    return OxmlElement(name)


def create_attribute(element, name, value):
    element.set(ns.qn(name), value)


def add_page_number(paragraph):
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    page_run = paragraph.add_run()
    t1 = create_element('w:t')
    create_attribute(t1, 'xml:space', 'preserve')
    t1.text = 'Page '
    page_run._r.append(t1)

    page_num_run = paragraph.add_run()

    fldChar1 = create_element('w:fldChar')
    create_attribute(fldChar1, 'w:fldCharType', 'begin')

    instrText = create_element('w:instrText')
    create_attribute(instrText, 'xml:space', 'preserve')
    instrText.text = "PAGE"

    fldChar2 = create_element('w:fldChar')
    create_attribute(fldChar2, 'w:fldCharType', 'end')

    page_num_run._r.append(fldChar1)
    page_num_run._r.append(instrText)
    page_num_run._r.append(fldChar2)

    of_run = paragraph.add_run()
    t2 = create_element('w:t')
    create_attribute(t2, 'xml:space', 'preserve')
    t2.text = ' of '
    of_run._r.append(t2)

    fldChar3 = create_element('w:fldChar')
    create_attribute(fldChar3, 'w:fldCharType', 'begin')

    instrText2 = create_element('w:instrText')
    create_attribute(instrText2, 'xml:space', 'preserve')
    instrText2.text = "NUMPAGES"

    fldChar4 = create_element('w:fldChar')
    create_attribute(fldChar4, 'w:fldCharType', 'end')

    num_pages_run = paragraph.add_run()
    num_pages_run._r.append(fldChar3)
    num_pages_run._r.append(instrText2)
    num_pages_run._r.append(fldChar4)

def create_table_ofdataframe(document,dataframe,heading=None):
    # add heading
    document.add_heading(heading, level=1)

    df_rows, df_cols = dataframe.shape
    table = document.add_table(rows=df_rows + 1, cols=df_cols)
    table.style = 'Table Grid'

    # Add column names to the first row of the table
    for col_num, column_name in enumerate(dataframe.columns):
        cell = table.cell(0, col_num)
        cell.text = column_name

    # Add data from the DataFrame to the table
    for row in range(df_rows):
        for col in range(df_cols):
            cell = table.cell(row + 1, col)
            cell.text = str(dataframe.iloc[row, col])
            # Adjust cell formatting (e.g., alignment)
            cell.paragraphs[0].alignment = 1  # Center alignment
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


# def save_documnet(input_path):
#     folder_input2 = st.text_input(':black','Testing Dataset',label_visibility='collapsed')
#     folder_input2=os.path.join(path, input_data_location,folder_input2)

 
